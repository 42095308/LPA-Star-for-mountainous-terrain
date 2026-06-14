from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


INPUT_DOCX = Path(r"C:\Users\42095\Desktop\小论文资料\稿件\初稿4.docx")
OUTPUT_DOCX = Path(r"C:\Users\42095\Desktop\小论文资料\稿件\初稿4_APA引用版.docx")


# 使用原文锚点定位，避免 Word 中标题、摘要或公式段落导致数字下标偏移。
CITATION_INSERTIONS = [
    ("无人机在山地物流、应急配送、巡检和灾害响应等任务中具有较高应用价值", "(Moshref-Javadi & Winkenbach, 2021; Tavana et al., 2017; Ramos & Vigo, 2023; Bhuiyan et al., 2024; Song et al., 2018; Aggarwal & Kumar, 2020)"),
    ("图搜索方法以A*、D*、LPA*和D*Lite等为代表", "(Hart et al., 1968; Stentz, 1994; Koenig & Likhachev, 2002; Koenig et al., 2004; Daniel et al., 2010; Aggarwal & Kumar, 2020)"),
    ("RRT、RRT*等采样方法能够在连续空间中快速探索可行区域", "(Karaman & Frazzoli, 2011; Fan et al., 2023; Zhang et al., 2023; Yu et al., 2023; Zhou et al., 2024)"),
    ("强化学习和深度强化学习能够通过环境交互学习路径策略", "(Zhao et al., 2018; Yu & Luo, 2023; Zhang et al., 2024)"),
    ("DEM、点云、三维栅格和体素模型能够表达真实地形", "(Sun et al., 2011; Ali et al., 2023; Huang & Savkin, 2021; Shao et al., 2025; Wang et al., 2024)"),
    ("动态重规划进一步增加了这一问题的复杂性", "(Koenig et al., 2004; Kim et al., 2019; Margraff et al., 2020; Volz & Graichen, 2019)"),
    ("人员暴露风险由OSM空间要素构建", "(Haklay & Weber, 2008; Mardani et al., 2019; Mohammed et al., 2021; Xu et al., 2023)"),
    ("本文将路径评价需求概括为时间效率、能耗经济性和风险规避三个方面", "(Wu et al., 2022; Salhi & Delavernhe, 2023; Zhou et al., 2024)"),
    ("本文基于DEM数据构建自适应安全飞行走廊", "(Ali et al., 2023; Huang & Savkin, 2021; Liu & Zou, 2025)"),
    ("在三层航线网络完成节点生成、边连接和安全校验后", "(Mardani et al., 2019; Salhi & Delavernhe, 2023; Wu et al., 2022; Zhou et al., 2024)"),
    ("为降低重规划成本，本文将区域扰动转化为航线图上的局部边阻断或边代价更新", "(Koenig & Likhachev, 2002; Koenig et al., 2004; Kim et al., 2019; Margraff et al., 2020)"),
    ("通过A*或LPA*搜索得到的结果是离散的航路点序列", "(Daniel et al., 2010; Nash et al., 2010; Chen et al., 2023; Nikolaiev & Novotarskyi, 2025)"),
    ("本文将完整方法记为MP，并设置MA、MF、MR和MV作为对比方法", "(Fan et al., 2023; Zhang et al., 2023; Yu et al., 2023; Yu & Luo, 2023; Zhou et al., 2024)"),
    ("本文面向真实山地环境下的无人机物流路径规划与动态重规划问题，提出了一种地形感知三层航线网络构建与事件驱动增量重规划方法", "(Bhuiyan et al., 2024; Mardani et al., 2019; Mohammed et al., 2021; Shao et al., 2025; Wang et al., 2024)"),
]


REFERENCES = [
    ("Aggarwal, S., & Kumar, N. (2020). Path planning techniques for unmanned aerial vehicles: A review, solutions, and challenges. Computer Communications, 149, 270-299.", "https://doi.org/10.1016/j.comcom.2019.10.014"),
    ("Ali, H., Xiong, G., Haider, M.H., Tamir, T.S., Dong, X., & Shen, Z. (2023). Feature selection-based decision model for UAV path planning on rough terrains. Expert Systems with Applications, 232, Article 120713.", "https://doi.org/10.1016/j.eswa.2023.120713"),
    ("Bhuiyan, T.H., Walker, V., Roni, M., & Ahmed, I. (2024). Aerial drone fleet deployment optimization with endogenous battery replacements for direct delivery of time-sensitive products. Expert Systems with Applications, 252, Article 124172.", "https://doi.org/10.1016/j.eswa.2024.124172"),
    ("Chen, W., Yang, Q., Diao, T., & Ren, S. (2023). B-Spline fusion line of sight algorithm for UAV path planning. Lecture Notes in Electrical Engineering, 503-512.", "https://doi.org/10.1007/978-981-19-6613-2_50"),
    ("Daniel, K., Nash, A., Koenig, S., & Felner, A. (2010). Theta*: Any-angle path planning on grids. Journal of Artificial Intelligence Research, 39, 533-579.", "https://doi.org/10.1613/jair.2994"),
    ("Fan, J., Chen, X., & Liang, X. (2023). UAV trajectory planning based on bi-directional APF-RRT* algorithm with goal-biased. Expert Systems with Applications, 213, Article 119137.", "https://doi.org/10.1016/j.eswa.2022.119137"),
    ("Haklay, M., & Weber, P. (2008). OpenStreetMap: User-generated street maps. IEEE Pervasive Computing, 7(4), 12-18.", "https://doi.org/10.1109/MPRV.2008.80"),
    ("Hart, P., Nilsson, N., & Raphael, B. (1968). A formal basis for the heuristic determination of minimum cost paths. IEEE Transactions on Systems Science and Cybernetics, 4(2), 100-107.", "https://doi.org/10.1109/TSSC.1968.300136"),
    ("Huang, H., & Savkin, A.V. (2021). Path planning for a solar-powered UAV inspecting mountain sites for safety and rescue. Energies, 14(7), Article 1968.", "https://doi.org/10.3390/en14071968"),
    ("Karaman, S., & Frazzoli, E. (2011). Sampling-based algorithms for optimal motion planning. The International Journal of Robotics Research, 30(7), 846-894.", "https://doi.org/10.1177/0278364911406761"),
    ("Kim, H., Jeong, J., Kim, N., & Kang, B. (2019). A study on 3D optimal path planning for quadcopter UAV based on D* Lite. In 2019 International Conference on Unmanned Aircraft Systems (ICUAS) (pp. 787-793).", "https://doi.org/10.1109/ICUAS.2019.8797815"),
    ("Koenig, S., & Likhachev, M. (2002). D* Lite. In Proceedings of the Eighteenth National Conference on Artificial Intelligence (AAAI) (pp. 476-483).", None),
    ("Koenig, S., Likhachev, M., & Furcy, D. (2004). Lifelong Planning A*. Artificial Intelligence, 155(1-2), 93-146.", "https://doi.org/10.1016/j.artint.2003.12.001"),
    ("Liu, J., & Zou, Y. (2025). Robust UAV path planning via safe flight corridor and penalty function. In 2025 40th Youth Academic Annual Conference of Chinese Association of Automation (YAC) (pp. 2138-2143).", "https://doi.org/10.1109/YAC66630.2025.11150073"),
    ("Mardani, A., Chiaberge, M., & Giaccone, P. (2019). Communication-aware UAV path planning. IEEE Access, 7, 52609-52621.", "https://doi.org/10.1109/ACCESS.2019.2911018"),
    ("Margraff, J., Stephant, J., & Labbani-Igbida, O. (2020). UAV 3D path and motion planning in unknown dynamic environments. In 2020 International Conference on Unmanned Aircraft Systems (ICUAS) (pp. 77-84).", "https://doi.org/10.1109/ICUAS48674.2020.9214057"),
    ("Mohammed, I., Collings, I.B., & Hanly, S.V. (2021). Line of sight probability prediction for UAV communication. In 2021 IEEE International Conference on Communications Workshops (ICC Workshops) (pp. 1-6).", "https://doi.org/10.1109/ICCWorkshops50388.2021.9473740"),
    ("Moshref-Javadi, M., & Winkenbach, M. (2021). Applications and research avenues for drone-based models in logistics: A classification and review. Expert Systems with Applications, 177, Article 114854.", "https://doi.org/10.1016/j.eswa.2021.114854"),
    ("Nash, A., Koenig, S., & Tovey, C. (2010). Lazy Theta*: Any-angle path planning and path length analysis in 3D. Proceedings of the AAAI Conference on Artificial Intelligence, 24(1), 147-154.", "https://doi.org/10.1609/aaai.v24i1.7566"),
    ("Nikolaiev, M., & Novotarskyi, M. (2025). An enhanced adaptive B-spline smoothing approach for UAV path planning. International Journal of Intelligent Systems and Applications, 17(4), 1-13.", "https://doi.org/10.5815/ijisa.2025.04.01"),
    ("Ramos, T.R.P., & Vigo, D. (2023). A new hybrid distribution paradigm: Integrating drones in medicines delivery. Expert Systems with Applications, 234, Article 120992.", "https://doi.org/10.1016/j.eswa.2023.120992"),
    ("Shao, Q., Mao, X., & Xu, W. (2025). Energy-aware UAV coverage planning in mountainous terrain via contour-aligned path generation. IEEE Robotics and Automation Letters, 10(12), 12373-12380.", "https://doi.org/10.1109/LRA.2025.3621932"),
    ("Song, B.D., Park, K., & Kim, J. (2018). Persistent UAV delivery logistics: MILP formulation and efficient heuristic. Computers & Industrial Engineering, 120, 418-428.", "https://doi.org/10.1016/j.cie.2018.05.013"),
    ("Stentz, A. (1994). Optimal and efficient path planning for partially-known environments. In Proceedings of the 1994 IEEE International Conference on Robotics and Automation (pp. 3310-3317).", "https://doi.org/10.1109/ROBOT.1994.351061"),
    ("Sun, T.Y., Huo, C.L., Tsai, S.J., Yu, Y.H., & Liu, C.C. (2011). Intelligent flight task algorithm for unmanned aerial vehicle. Expert Systems with Applications, 38(8), 10036-10048.", "https://doi.org/10.1016/j.eswa.2011.02.013"),
    ("Tavana, M., Khalili-Damghani, K., Santos-Arteaga, F.J., & Zandi, M.H. (2017). Drone shipping versus truck delivery in a cross-docking system with multiple fleets and products. Expert Systems with Applications, 72, 93-107.", "https://doi.org/10.1016/j.eswa.2016.12.014"),
    ("Volz, A., & Graichen, K. (2019). A predictive path-following controller for continuous replanning with dynamic roadmaps. IEEE Robotics and Automation Letters, 4(4), 3963-3970.", "https://doi.org/10.1109/LRA.2019.2929990"),
    ("Wang, W., Li, X., & Tian, J. (2024). UAV formation path planning for mountainous forest terrain utilizing an artificial rabbit optimizer incorporating reinforcement learning and thermal conduction search strategies. Advanced Engineering Informatics, 62, Article 102947.", "https://doi.org/10.1016/j.aei.2024.102947"),
    ("Wu, M., Chen, W., & Tian, X. (2022). Optimal energy consumption path planning for quadrotor UAV transmission tower inspection based on simulated annealing algorithm. Energies, 15(21), Article 8036.", "https://doi.org/10.3390/en15218036"),
    ("Xu, L., Cao, X., Du, W., & Li, Y. (2023). Cooperative path planning optimization for multiple UAVs with communication constraints. Knowledge-Based Systems, 260, Article 110164.", "https://doi.org/10.1016/j.knosys.2022.110164"),
    ("Yu, X., Jiang, N., Wang, X., & Li, M. (2023). A hybrid algorithm based on grey wolf optimizer and differential evolution for UAV path planning. Expert Systems with Applications, 215, Article 119327.", "https://doi.org/10.1016/j.eswa.2022.119327"),
    ("Yu, X., & Luo, W. (2023). Reinforcement learning-based multi-strategy cuckoo search algorithm for 3D UAV path planning. Expert Systems with Applications, 223, Article 119910.", "https://doi.org/10.1016/j.eswa.2023.119910"),
    ("Zhao, Y., Zheng, Z., & Liu, Y. (2018). Survey on computational-intelligence-based UAV path planning. Knowledge-Based Systems, 158, 54-64.", "https://doi.org/10.1016/j.knosys.2018.05.033"),
    ("Zhang, C., Zhou, W., Qin, W., & Tang, W. (2023). A novel UAV path planning approach: Heuristic crossing search and rescue optimization algorithm. Expert Systems with Applications, 215, Article 119243.", "https://doi.org/10.1016/j.eswa.2022.119243"),
    ("Zhang, W., Peng, C., Yuan, Y., Cui, J., & Qi, L. (2024). A novel multi-objective evolutionary algorithm with a two-fold constraint-handling mechanism for multiple UAV path planning. Expert Systems with Applications, 238, Article 121862.", "https://doi.org/10.1016/j.eswa.2023.121862"),
    ("Zhou, X., Tang, Z., Wang, N., Yang, C., & Huang, T. (2024). A novel state transition algorithm with adaptive fuzzy penalty for multi-constraint UAV path planning. Expert Systems with Applications, 248, Article 123481.", "https://doi.org/10.1016/j.eswa.2024.123481"),
]


def add_hyperlink(paragraph, text, url):
    """向段落添加 DOI 超链接。"""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    new_run.append(r_pr)

    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def append_citation_before_terminal_punctuation(paragraph, citation):
    """尽量把引用放在段落末尾标点之前，减少对原文格式的扰动。"""
    if citation in paragraph.text:
        return

    terminal_marks = "。.!！?？；;"
    for run in reversed(paragraph.runs):
        if not run.text:
            continue
        stripped = run.text.rstrip()
        trailing_space = run.text[len(stripped) :]
        if stripped and stripped[-1] in terminal_marks:
            mark = stripped[-1]
            run.text = stripped[:-1] + trailing_space
            paragraph.add_run(f" {citation}{mark}")
            return
        paragraph.add_run(f" {citation}")
        return

    paragraph.add_run(citation)


def add_reference_paragraph(document, ref_text, doi_url):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.5)
    paragraph.paragraph_format.space_after = 0
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.add_run(ref_text)
    if doi_url:
        paragraph.add_run(" ")
        add_hyperlink(paragraph, doi_url, doi_url)
    return paragraph


def find_paragraph_by_anchor(document, anchor):
    """根据原文锚点查找段落。"""
    for paragraph in document.paragraphs:
        if anchor in paragraph.text:
            return paragraph
    raise ValueError(f"未找到原文锚点: {anchor}")


def prepare_reference_heading(document):
    """复用原稿末尾的文献引用标题；若不存在则新增 References 标题。"""
    for paragraph in document.paragraphs:
        if paragraph.text.strip() in {"文献引用", "References"}:
            paragraph.text = "References"
            return
    document.add_page_break()
    document.add_heading("References", level=1)


def main():
    if not INPUT_DOCX.exists():
        raise FileNotFoundError(f"找不到输入文件: {INPUT_DOCX}")

    document = Document(str(INPUT_DOCX))

    for anchor, citation in CITATION_INSERTIONS:
        paragraph = find_paragraph_by_anchor(document, anchor)
        append_citation_before_terminal_punctuation(paragraph, citation)

    prepare_reference_heading(document)
    for ref_text, doi_url in REFERENCES:
        add_reference_paragraph(document, ref_text, doi_url)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(OUTPUT_DOCX))
    print(f"已生成: {OUTPUT_DOCX}")
    print(f"插入文内引用位置数: {len(CITATION_INSERTIONS)}")
    print(f"参考文献条目数: {len(REFERENCES)}")


if __name__ == "__main__":
    main()
