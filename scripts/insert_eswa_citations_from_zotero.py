from __future__ import annotations

import re
import shutil
import sqlite3
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


INPUT_DOCX = Path(r"C:\Users\42095\Desktop\小论文资料\稿件\初稿4.docx")
OUTPUT_DOCX = Path(r"C:\Users\42095\Desktop\小论文资料\稿件\初稿4_Zotero_ESWA引用版.docx")
ZOTERO_DB = Path(r"C:\Users\42095\Zotero\zotero.sqlite")
DB_COPY = Path(r"C:\programming\code\article_python\tmp\zotero\zotero.sqlite")


@dataclass
class ZoteroItem:
    item_id: int
    item_type: str
    fields: dict[str, str] = field(default_factory=dict)
    creators: list[str] = field(default_factory=list)

    @property
    def title(self) -> str:
        return self.fields.get("title", "").strip()

    @property
    def year(self) -> str:
        date = self.fields.get("date", "")
        match = re.search(r"(19|20)\d{2}", date)
        return match.group(0) if match else "n.d."

    @property
    def doi(self) -> str:
        return self.fields.get("DOI", "").strip()

    @property
    def doi_url(self) -> str:
        return f"https://doi.org/{self.doi}" if self.doi else ""

    @property
    def url(self) -> str:
        return self.fields.get("url", "").strip()

    @property
    def sort_key(self):
        first_author = self.creators[0].split()[-1].lower() if self.creators else ""
        return first_author, self.year, self.title.lower()


CITATION_INSERTIONS = [
    ("无人机在山地物流、应急配送、巡检和灾害响应等任务中具有较高应用价值", "(Moshref-Javadi & Winkenbach, 2021; Tavana et al., 2017; Ramos & Vigo, 2023; Bhuiyan et al., 2024; Song et al., 2018; Aggarwal & Kumar, 2020)"),
    ("图搜索方法以A*、D*、LPA*和D*Lite等为代表", "(Hart et al., 1968; Stentz, 1994; Koenig & Likhachev, 2002; Koenig et al., 2004; Daniel et al., 2010; Aggarwal & Kumar, 2020)"),
    ("RRT、RRT*等采样方法能够在连续空间中快速探索可行区域", "(Karaman & Frazzoli, 2011; Fan et al., 2023; Zhang et al., 2023; Yu et al., 2023; Zhou et al., 2024)"),
    ("强化学习和深度强化学习能够通过环境交互学习路径策略", "(Zhao et al., 2018; Yu & Luo, 2023; Zhang et al., 2024)"),
    ("DEM、点云、三维栅格和体素模型能够表达真实地形", "(Sun et al., 2011; Ali et al., 2023; Huang & Savkin, 2021; Shao et al., 2025; Wang et al., 2024)"),
    ("动态重规划进一步增加了这一问题的复杂性", "(Koenig et al., 2004; Kim et al., 2019; Margraff et al., 2020; Volz & Graichen, 2019)"),
    ("人员暴露风险由OSM空间要素构建", "(Haklay & Weber, 2008; Mardani et al., 2019; Mohammed et al., 2021; Xu et al., 2023)"),
    ("本文将路径评价需求概括为时间效率、能耗经济性和风险规避三个方面", "(Wu et al., 2022; Salhi & Delavernhe, 2023; Zhou et al., 2024)"),
    ("本文基于DEM数据构建自适应安全飞行走廊", "(Ali et al., 2023; Huang & Savkin, 2021; Liu & Zou, 2025)"),
    ("若在完整三维空间中均匀生成候选节点", "(Hu et al., 2025; Ali et al., 2023; Shao et al., 2025)"),
    ("在三层航线网络完成节点生成、边连接和安全校验后", "(Mardani et al., 2019; Salhi & Delavernhe, 2023; Wu et al., 2022; Zhou et al., 2024)"),
    ("为降低重规划成本，本文将区域扰动转化为航线图上的局部边阻断或边代价更新", "(Koenig & Likhachev, 2002; Koenig et al., 2004; Kim et al., 2019; Margraff et al., 2020)"),
    ("通过A*或LPA*搜索得到的结果是离散的航路点序列", "(Daniel et al., 2010; Nash et al., 2010; Chen et al., 2023; Nikolaiev & Novotarskyi, 2025)"),
    ("本文将完整方法记为MP，并设置MA、MF、MR和MV作为对比方法", "(Fan et al., 2023; Zhang et al., 2023; Yu et al., 2023; Yu & Luo, 2023; Zhou et al., 2024)"),
    ("本文面向真实山地环境下的无人机物流路径规划与动态重规划问题，提出了一种地形感知三层航线网络构建与事件驱动增量重规划方法", "(Bhuiyan et al., 2024; Mardani et al., 2019; Mohammed et al., 2021; Shao et al., 2025; Wang et al., 2024)"),
]


def copy_zotero_db() -> None:
    DB_COPY.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(ZOTERO_DB, DB_COPY)


def read_zotero_items() -> list[ZoteroItem]:
    con = sqlite3.connect(DB_COPY)
    con.row_factory = sqlite3.Row
    cur = con.cursor()
    rows = cur.execute(
        """
        select i.itemID, it.typeName, f.fieldName, v.value
        from items i
        join itemTypes it on i.itemTypeID = it.itemTypeID
        left join itemData d on i.itemID = d.itemID
        left join fields f on d.fieldID = f.fieldID
        left join itemDataValues v on d.valueID = v.valueID
        where i.itemID not in (select itemID from deletedItems)
        order by i.itemID
        """
    )
    items: dict[int, ZoteroItem] = {}
    for row in rows:
        item = items.setdefault(row["itemID"], ZoteroItem(row["itemID"], row["typeName"]))
        if row["fieldName"]:
            item.fields[row["fieldName"]] = row["value"]

    creator_rows = cur.execute(
        """
        select ic.itemID, ic.orderIndex, c.firstName, c.lastName, c.fieldMode
        from itemCreators ic
        join creators c on ic.creatorID = c.creatorID
        order by ic.itemID, ic.orderIndex
        """
    )
    for row in creator_rows:
        if row["itemID"] not in items:
            continue
        if row["fieldMode"] == 1:
            name = row["lastName"] or ""
        else:
            name = f"{row['firstName'] or ''} {row['lastName'] or ''}".strip()
        if name:
            items[row["itemID"]].creators.append(name)
    con.close()

    return [
        item
        for item in items.values()
        if item.item_type not in {"attachment", "annotation"} and item.title
    ]


def family_and_initials(name: str) -> str:
    parts = name.split()
    if not parts:
        return ""
    family = parts[-1]
    given = parts[:-1]
    initials = "".join(part[0].upper() + "." for part in given if part)
    return f"{family}, {initials}" if initials else family


def format_author_list(creators: list[str]) -> str:
    formatted = [family_and_initials(name) for name in creators]
    formatted = [name for name in formatted if name]
    if not formatted:
        return ""
    if len(formatted) == 1:
        return formatted[0]
    return ", ".join(formatted[:-1]) + ", & " + formatted[-1]


def normalize_pages(pages: str) -> str:
    return pages.replace("–", "-").replace("—", "-").strip()


def container_title(item: ZoteroItem) -> str:
    fields = item.fields
    return (
        fields.get("publicationTitle")
        or fields.get("proceedingsTitle")
        or fields.get("conferenceName")
        or fields.get("bookTitle")
        or ""
    ).strip()


def format_reference(item: ZoteroItem) -> tuple[str, str]:
    authors = format_author_list(item.creators)
    title = item.title.rstrip(".")
    year = item.year
    fields = item.fields
    volume = fields.get("volume", "").strip()
    issue = fields.get("issue", "").strip()
    pages = normalize_pages(fields.get("pages", ""))
    container = container_title(item)

    if item.item_type in {"conferencePaper"}:
        ref = f"{authors} ({year}). {title}. In {container}"
        if pages:
            ref += f" (pp. {pages})"
        ref += "."
    elif item.item_type in {"bookSection"}:
        ref = f"{authors} ({year}). {title}. In {container}"
        if pages:
            ref += f" (pp. {pages})"
        ref += "."
    else:
        ref = f"{authors} ({year}). {title}. {container}"
        if volume:
            ref += f", {volume}"
            if issue:
                ref += f"({issue})"
        if pages:
            if pages.isdigit() and len(pages) >= 5:
                ref += f", Article {pages}"
            else:
                ref += f", {pages}"
        ref += "."

    link = item.doi_url or item.url
    return ref, link


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def append_citation_before_terminal_punctuation(paragraph, citation: str) -> None:
    if citation in paragraph.text:
        return
    terminal_marks = "。.!！?？；;"
    for run in reversed(paragraph.runs):
        if not run.text:
            continue
        stripped = run.text.rstrip()
        trailing = run.text[len(stripped):]
        if stripped and stripped[-1] in terminal_marks:
            mark = stripped[-1]
            run.text = stripped[:-1] + trailing
            paragraph.add_run(f" {citation}{mark}")
            return
        paragraph.add_run(f" {citation}")
        return
    paragraph.add_run(citation)


def find_paragraph_by_anchor(document: Document, anchor: str):
    for paragraph in document.paragraphs:
        if anchor in paragraph.text:
            return paragraph
    raise ValueError(f"未找到原文锚点: {anchor}")


def prepare_reference_heading(document: Document) -> None:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() in {"文献引用", "References"}:
            paragraph.text = "References"
            return
    document.add_page_break()
    document.add_heading("References", level=1)


def add_reference_paragraph(document: Document, ref_text: str, link: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.5)
    paragraph.paragraph_format.space_after = 0
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.add_run(ref_text)
    if link:
        paragraph.add_run(" ")
        add_hyperlink(paragraph, link, link)


def main() -> None:
    if not INPUT_DOCX.exists():
        raise FileNotFoundError(f"找不到初稿: {INPUT_DOCX}")
    if not ZOTERO_DB.exists():
        raise FileNotFoundError(f"找不到 Zotero 数据库: {ZOTERO_DB}")

    copy_zotero_db()
    items = sorted(read_zotero_items(), key=lambda item: item.sort_key)
    document = Document(str(INPUT_DOCX))

    for anchor, citation in CITATION_INSERTIONS:
        append_citation_before_terminal_punctuation(find_paragraph_by_anchor(document, anchor), citation)

    prepare_reference_heading(document)
    for item in items:
        ref_text, link = format_reference(item)
        add_reference_paragraph(document, ref_text, link)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(OUTPUT_DOCX))
    doi_count = sum(1 for item in items if item.doi)
    no_doi = [item.title for item in items if not item.doi]
    print(f"已生成: {OUTPUT_DOCX}")
    print(f"Zotero文献条目数: {len(items)}")
    print(f"带DOI条目数: {doi_count}")
    print(f"无DOI条目: {no_doi}")


if __name__ == "__main__":
    main()
