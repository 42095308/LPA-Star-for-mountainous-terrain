from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


INPUT_DOCX = Path(r"C:\Users\42095\Desktop\小论文资料\稿件\初稿5-已引用版.docx")
OUTPUT_DOCX = Path(r"C:\Users\42095\Desktop\小论文资料\稿件\初稿5-引用重构版.docx")
REPORT_PATH = Path(r"C:\programming\code\article_python\output\literature\初稿5_引用重构审计.md")


@dataclass(frozen=True)
class Reference:
    key: str
    text: str
    link: str | None
    year: int
    venue: str
    note: str


# 只替换完整引用组，避免误伤正文中的算法名称或作者姓名。
CITATION_REPLACEMENTS = [
    (
        "(Moshref-Javadi & Winkenbach, 2021; Tavana et al., 2017; Ramos & Vigo, 2023; Bhuiyan et al., 2024; Song et al., 2018; Aggarwal & Kumar, 2020)",
        "(Moshref-Javadi & Winkenbach, 2021; Ramos & Vigo, 2023; Bhuiyan et al., 2024; Mishra & Tiwari, 2024; Yılmaz et al., 2024; Aggarwal & Kumar, 2020)",
        "引言应用背景：移除 2017、2018 年无人机配送旧文献，补入 2024 年 ESWA 物流配送文献。",
    ),
    (
        "(Hart et al., 1968; Stentz, 1994; Koenig & Likhachev, 2002; Koenig et al., 2004; Daniel et al., 2010; Aggarwal & Kumar, 2020)",
        "(Hart et al., 1968; Stentz, 1994; Koenig & Likhachev, 2002; Koenig et al., 2004; Aggarwal & Kumar, 2020)",
        "图搜索综述：保留 A*、D*、D* Lite、LPA* 经典源文献，移除非必要的 Theta* 旧文献。",
    ),
    (
        "(Zhao et al., 2018; Yu & Luo, 2023; Zhang et al., 2024)",
        "(Aggarwal & Kumar, 2020; Ghambari et al., 2024; Yu & Luo, 2023; Zhang et al., 2024)",
        "学习类方法综述：用 2020 年后综述替换 2018 年计算智能综述。",
    ),
    (
        "(Sun et al., 2011; Ali et al., 2023; Huang & Savkin, 2021; Shao et al., 2025; Wang et al., 2024)",
        "(Ali et al., 2023; Huang & Savkin, 2021; Hu et al., 2025; Shao et al., 2025; Wang et al., 2024)",
        "三维环境表达：移除 2011 年旧 ESWA 任务算法文献，改用 2020 年后的地形、分层图和山地规划文献。",
    ),
    (
        "(Koenig et al., 2004; Kim et al., 2019; Margraff et al., 2020; Volz & Graichen, 2019)",
        "(Koenig et al., 2004; Margraff et al., 2020; Simplicio & Pereira, 2025)",
        "动态重规划背景：保留 LPA* 经典源文献，移除 2019 年应用型旧文献，补入 2025 年重规划文献。",
    ),
    (
        "(Haklay & Weber, 2008; Mardani et al., 2019; Mohammed et al., 2021; Xu et al., 2023)",
        "(Bertolotto et al., 2020; Cichociński, 2021; Mohammed et al., 2021; Xu et al., 2023)",
        "环境输入定义：用 2020 年后的 OSM 和开放空间数据论文替换 2008 年 OSM 旧文献，同时移除 2019 年通信感知路径规划文献。",
    ),
    (
        "(Mardani et al., 2019; Salhi & Delavernhe, 2023; Wu et al., 2022; Zhou et al., 2024)",
        "(Mohammed et al., 2021; Salhi & Delavernhe, 2023; Wu et al., 2022; Xu et al., 2023; Zhou et al., 2024)",
        "边代价建模：移除 2019 年通信感知路径规划文献，改由 2021 年后通信约束和多目标代价文献支撑。",
    ),
    (
        "(Koenig & Likhachev, 2002; Koenig et al., 2004; Kim et al., 2019; Margraff et al., 2020)",
        "(Koenig & Likhachev, 2002; Koenig et al., 2004; Margraff et al., 2020; Simplicio & Pereira, 2025)",
        "增量重规划方法：保留 D* Lite 与 LPA* 经典源文献，移除 2019 年 UAV 应用文献。",
    ),
    (
        "(Daniel et al., 2010; Nash et al., 2010; Chen et al., 2023; Nikolaiev & Novotarskyi, 2025)",
        "(Chen et al., 2023; Nikolaiev & Novotarskyi, 2025)",
        "路径平滑后处理：移除 2010 年 any-angle 旧文献，保留 2023、2025 年 B 样条和 UAV 平滑处理文献。",
    ),
    (
        "(Bhuiyan et al., 2024; Mardani et al., 2019; Mohammed et al., 2021; Shao et al., 2025; Wang et al., 2024)",
        "(Bhuiyan et al., 2024; Mohammed et al., 2021; Shao et al., 2025; Wang et al., 2024)",
        "结论回扣：移除 2019 年旧文献，保留 2021 年后的物流、通信、山地路径规划支撑。",
    ),
]


REFERENCES = [
    Reference(
        "Aggarwal2020",
        "Aggarwal, S., & Kumar, N. (2020). Path planning techniques for unmanned aerial vehicles: A review, solutions, and challenges. Computer Communications, 149, 270-299.",
        "https://doi.org/10.1016/j.comcom.2019.10.014",
        2020,
        "Computer Communications",
        "2020 年后 UAV 路径规划综述，替代部分旧综述。",
    ),
    Reference(
        "Ali2023",
        "Ali, H., Xiong, G., Haider, M.H., Tamir, T.S., Dong, X., & Shen, Z. (2023). Feature selection-based decision model for UAV path planning on rough terrains. Expert Systems with Applications, 232, Article 120713.",
        "https://doi.org/10.1016/j.eswa.2023.120713",
        2023,
        "Expert Systems with Applications",
        "ESWA 粗糙地形 UAV 路径规划。",
    ),
    Reference(
        "Bertolotto2020",
        "Bertolotto, M., McArdle, G., & Schoen-Phelan, B. (2020). Volunteered and crowdsourced geographic information: The OpenStreetMap project. Journal of Spatial Information Science, 20.",
        "https://doi.org/10.5311/josis.2020.20.659",
        2020,
        "Journal of Spatial Information Science",
        "替代 2008 年 OSM 旧文献。",
    ),
    Reference(
        "Bhuiyan2024",
        "Bhuiyan, T.H., Walker, V., Roni, M., & Ahmed, I. (2024). Aerial drone fleet deployment optimization with endogenous battery replacements for direct delivery of time-sensitive products. Expert Systems with Applications, 252, Article 124172.",
        "https://doi.org/10.1016/j.eswa.2024.124172",
        2024,
        "Expert Systems with Applications",
        "ESWA 无人机配送优化。",
    ),
    Reference(
        "Chen2023",
        "Chen, W., Yang, Q., Diao, T., & Ren, S. (2023). B-spline fusion line of sight algorithm for UAV path planning. In Lecture Notes in Electrical Engineering (pp. 503-512).",
        "https://doi.org/10.1007/978-981-19-6613-2_50",
        2023,
        "Lecture Notes in Electrical Engineering",
        "路径平滑和视距处理。",
    ),
    Reference(
        "Cichocinski2021",
        "Cichociński, P. (2021). A study on the usability of open spatial data for road network-based analysis: Using OpenStreetMap as an example. Geoinformatica Polonica, 20, 89-96.",
        "https://doi.org/10.4467/21995923GP.21.007.14978",
        2021,
        "Geoinformatica Polonica",
        "OSM 道路网络和开放空间数据支撑。",
    ),
    Reference(
        "Fan2023",
        "Fan, J., Chen, X., & Liang, X. (2023). UAV trajectory planning based on bi-directional APF-RRT* algorithm with goal-biased. Expert Systems with Applications, 213, Article 119137.",
        "https://doi.org/10.1016/j.eswa.2022.119137",
        2023,
        "Expert Systems with Applications",
        "ESWA UAV 轨迹规划对比文献。",
    ),
    Reference(
        "Ghambari2024",
        "Ghambari, S., Golabi, M., Jourdan, L., Lepagnot, J., & Idoumghar, L. (2024). UAV path planning techniques: A survey. RAIRO - Operations Research, 58(4), 2951-2989.",
        "https://doi.org/10.1051/ro/2024073",
        2024,
        "RAIRO - Operations Research",
        "2024 年 UAV 路径规划综述。",
    ),
    Reference(
        "Hart1968",
        "Hart, P., Nilsson, N., & Raphael, B. (1968). A formal basis for the heuristic determination of minimum cost paths. IEEE Transactions on Systems Science and Cybernetics, 4(2), 100-107.",
        "https://doi.org/10.1109/TSSC.1968.300136",
        1968,
        "IEEE Transactions on Systems Science and Cybernetics",
        "A* 经典源文献，按用户规则保留。",
    ),
    Reference(
        "Hu2025",
        "Hu, X., Yang, C., Zhou, J., Zhang, Y., & Ma, Y. (2025). Research on 3D layered visibility graph route network model and multi-objective path planning for UAVs in complex urban environments. Aerospace Science and Technology, 159, Article 109947.",
        "https://doi.org/10.1016/j.ast.2025.109947",
        2025,
        "Aerospace Science and Technology",
        "2025 年三维分层图航线网络。",
    ),
    Reference(
        "Huang2021",
        "Huang, H., & Savkin, A.V. (2021). Path planning for a solar-powered UAV inspecting mountain sites for safety and rescue. Energies, 14(7), Article 1968.",
        "https://doi.org/10.3390/en14071968",
        2021,
        "Energies",
        "山地站点巡检路径规划。",
    ),
    Reference(
        "Karaman2011",
        "Karaman, S., & Frazzoli, E. (2011). Sampling-based algorithms for optimal motion planning. The International Journal of Robotics Research, 30(7), 846-894.",
        "https://doi.org/10.1177/0278364911406761",
        2011,
        "The International Journal of Robotics Research",
        "RRT* 经典源文献，按用户规则保留。",
    ),
    Reference(
        "Koenig2002",
        "Koenig, S., & Likhachev, M. (2002). D* Lite. In Proceedings of the 18th AAAI Conference on Artificial Intelligence (AAAI '02) (pp. 476-483).",
        "https://aaai.org/papers/00476-aaai02-072-d-lite/",
        2002,
        "AAAI Conference on Artificial Intelligence",
        "D* Lite 经典源文献，无 DOI，保留 URL。",
    ),
    Reference(
        "Koenig2004",
        "Koenig, S., Likhachev, M., & Furcy, D. (2004). Lifelong Planning A*. Artificial Intelligence, 155(1-2), 93-146.",
        "https://doi.org/10.1016/j.artint.2003.12.001",
        2004,
        "Artificial Intelligence",
        "LPA* 经典源文献，按用户规则保留。",
    ),
    Reference(
        "Liu2025",
        "Liu, J., & Zou, Y. (2025). Robust UAV path planning via safe flight corridor and penalty function. In 2025 40th Youth Academic Annual Conference of Chinese Association of Automation (YAC) (pp. 2138-2143).",
        "https://doi.org/10.1109/YAC66630.2025.11150073",
        2025,
        "YAC",
        "安全飞行走廊相关。",
    ),
    Reference(
        "Margraff2020",
        "Margraff, J., Stephant, J., & Labbani-Igbida, O. (2020). UAV 3D path and motion planning in unknown dynamic environments. In 2020 International Conference on Unmanned Aircraft Systems (ICUAS) (pp. 77-84).",
        "https://doi.org/10.1109/ICUAS48674.2020.9214057",
        2020,
        "ICUAS",
        "2020 年后动态环境 UAV 路径规划。",
    ),
    Reference(
        "Mishra2024",
        "Mishra, D., & Tiwari, M.K. (2024). Integrated truck drone delivery services with an optimal charging stations. Expert Systems with Applications, 254, Article 124254.",
        "https://doi.org/10.1016/j.eswa.2024.124254",
        2024,
        "Expert Systems with Applications",
        "ESWA 2024 卡车、无人机配送服务。",
    ),
    Reference(
        "Mohammed2021",
        "Mohammed, I., Collings, I.B., & Hanly, S.V. (2021). Line of sight probability prediction for UAV communication. In 2021 IEEE International Conference on Communications Workshops (ICC Workshops) (pp. 1-6).",
        "https://doi.org/10.1109/ICCWorkshops50388.2021.9473740",
        2021,
        "ICC Workshops",
        "UAV 通信视距概率预测。",
    ),
    Reference(
        "Moshref2021",
        "Moshref-Javadi, M., & Winkenbach, M. (2021). Applications and research avenues for drone-based models in logistics: A classification and review. Expert Systems with Applications, 177, Article 114854.",
        "https://doi.org/10.1016/j.eswa.2021.114854",
        2021,
        "Expert Systems with Applications",
        "ESWA 无人机物流综述。",
    ),
    Reference(
        "Nikolaiev2025",
        "Nikolaiev, M., & Novotarskyi, M. (2025). An enhanced adaptive B-spline smoothing approach for UAV path planning. International Journal of Intelligent Systems and Applications, 17(4), 1-13.",
        "https://doi.org/10.5815/ijisa.2025.04.01",
        2025,
        "International Journal of Intelligent Systems and Applications",
        "2025 年 UAV B 样条平滑。",
    ),
    Reference(
        "Ramos2023",
        "Ramos, T.R.P., & Vigo, D. (2023). A new hybrid distribution paradigm: Integrating drones in medicines delivery. Expert Systems with Applications, 234, Article 120992.",
        "https://doi.org/10.1016/j.eswa.2023.120992",
        2023,
        "Expert Systems with Applications",
        "ESWA 药品无人机配送。",
    ),
    Reference(
        "Salhi2023",
        "Salhi, M., & Delavernhe, F. (2023). Multiobjective UAV path planning: Connectivity quality and energy consumption. In GLOBECOM 2023 - 2023 IEEE Global Communications Conference (pp. 746-751).",
        "https://doi.org/10.1109/GLOBECOM54140.2023.10437552",
        2023,
        "GLOBECOM",
        "通信质量和能耗多目标代价。",
    ),
    Reference(
        "Shao2025",
        "Shao, Q., Mao, X., & Xu, W. (2025). Energy-aware UAV coverage planning in mountainous terrain via contour-aligned path generation. IEEE Robotics and Automation Letters, 10(12), 12373-12380.",
        "https://doi.org/10.1109/LRA.2025.3621932",
        2025,
        "IEEE Robotics and Automation Letters",
        "山地地形能耗感知 UAV 路径规划。",
    ),
    Reference(
        "Simplicio2025",
        "Simplicio, P.V.G., & Pereira, G.A.S. (2025). Multi-resolution UAV path replanning for inspection of tailings dams. In 2025 International Conference on Unmanned Aircraft Systems (ICUAS) (pp. 256-263).",
        "https://doi.org/10.1109/ICUAS65942.2025.11007892",
        2025,
        "ICUAS",
        "2025 年 UAV 路径重规划。",
    ),
    Reference(
        "Stentz1994",
        "Stentz, A. (1994). Optimal and efficient path planning for partially-known environments. In Proceedings of the 1994 IEEE International Conference on Robotics and Automation (pp. 3310-3317).",
        "https://doi.org/10.1109/ROBOT.1994.351061",
        1994,
        "IEEE International Conference on Robotics and Automation",
        "D* 经典源文献，按用户规则保留。",
    ),
    Reference(
        "Wang2024",
        "Wang, W., Li, X., & Tian, J. (2024). UAV formation path planning for mountainous forest terrain utilizing an artificial rabbit optimizer incorporating reinforcement learning and thermal conduction search strategies. Advanced Engineering Informatics, 62, Article 102947.",
        "https://doi.org/10.1016/j.aei.2024.102947",
        2024,
        "Advanced Engineering Informatics",
        "山地森林地形 UAV 编队路径规划。",
    ),
    Reference(
        "Wu2022",
        "Wu, M., Chen, W., & Tian, X. (2022). Optimal energy consumption path planning for quadrotor UAV transmission tower inspection based on simulated annealing algorithm. Energies, 15(21), Article 8036.",
        "https://doi.org/10.3390/en15218036",
        2022,
        "Energies",
        "能耗路径代价。",
    ),
    Reference(
        "Xu2023",
        "Xu, L., Cao, X., Du, W., & Li, Y. (2023). Cooperative path planning optimization for multiple UAVs with communication constraints. Knowledge-Based Systems, 260, Article 110164.",
        "https://doi.org/10.1016/j.knosys.2022.110164",
        2023,
        "Knowledge-Based Systems",
        "通信约束多 UAV 路径规划。",
    ),
    Reference(
        "Yilmaz2024",
        "Yılmaz, C., Cengiz, E., & Kahraman, H.T. (2024). A new evolutionary optimization algorithm with hybrid guidance mechanism for truck-multi drone delivery system. Expert Systems with Applications, 245, Article 123115.",
        "https://doi.org/10.1016/j.eswa.2023.123115",
        2024,
        "Expert Systems with Applications",
        "ESWA 2024 卡车、多无人机配送优化。",
    ),
    Reference(
        "Yu2023",
        "Yu, X., Jiang, N., Wang, X., & Li, M. (2023). A hybrid algorithm based on grey wolf optimizer and differential evolution for UAV path planning. Expert Systems with Applications, 215, Article 119327.",
        "https://doi.org/10.1016/j.eswa.2022.119327",
        2023,
        "Expert Systems with Applications",
        "ESWA UAV 路径规划智能优化。",
    ),
    Reference(
        "YuLuo2023",
        "Yu, X., & Luo, W. (2023). Reinforcement learning-based multi-strategy cuckoo search algorithm for 3D UAV path planning. Expert Systems with Applications, 223, Article 119910.",
        "https://doi.org/10.1016/j.eswa.2023.119910",
        2023,
        "Expert Systems with Applications",
        "ESWA 强化学习与 3D UAV 路径规划。",
    ),
    Reference(
        "ZhangC2023",
        "Zhang, C., Zhou, W., Qin, W., & Tang, W. (2023). A novel UAV path planning approach: Heuristic crossing search and rescue optimization algorithm. Expert Systems with Applications, 215, Article 119243.",
        "https://doi.org/10.1016/j.eswa.2022.119243",
        2023,
        "Expert Systems with Applications",
        "ESWA UAV 路径规划对比文献。",
    ),
    Reference(
        "ZhangW2024",
        "Zhang, W., Peng, C., Yuan, Y., Cui, J., & Qi, L. (2024). A novel multi-objective evolutionary algorithm with a two-fold constraint-handling mechanism for multiple UAV path planning. Expert Systems with Applications, 238, Article 121862.",
        "https://doi.org/10.1016/j.eswa.2023.121862",
        2024,
        "Expert Systems with Applications",
        "ESWA 多目标多 UAV 路径规划。",
    ),
    Reference(
        "Zhou2024",
        "Zhou, X., Tang, Z., Wang, N., Yang, C., & Huang, T. (2024). A novel state transition algorithm with adaptive fuzzy penalty for multi-constraint UAV path planning. Expert Systems with Applications, 248, Article 123481.",
        "https://doi.org/10.1016/j.eswa.2024.123481",
        2024,
        "Expert Systems with Applications",
        "ESWA 多约束 UAV 路径规划。",
    ),
]


REMOVED_REFERENCES = [
    "Daniel et al. (2010)：Theta*，不属于本文必须保留的经典源算法，且后处理段落已有 2023、2025 年文献。",
    "Haklay & Weber (2008)：OSM 早期介绍文献，替换为 2020、2021 年开放空间数据和 OSM 文献。",
    "Kim et al. (2019)：D* Lite 的 UAV 应用文献，保留 D* Lite 源文献即可。",
    "Mardani et al. (2019)：通信感知 UAV 路径规划旧文献，替换为 2021 年后通信视距和通信约束文献。",
    "Nash et al. (2010)：Lazy Theta* 旧文献，当前段落不是 any-angle 算法综述，移除。",
    "Song et al. (2018)：无人机配送旧文献，替换为 2024 年 ESWA 物流配送文献。",
    "Sun et al. (2011)：旧 ESWA UAV flight task 文献，替换为 2023 年后地形、分层图和山地规划文献。",
    "Tavana et al. (2017)：旧 ESWA 卡车、无人机配送文献，替换为 2024 年 ESWA 同类文献。",
    "Volz & Graichen (2019)：连续重规划旧应用文献，替换为 2025 年 UAV 重规划文献。",
    "Zhao et al. (2018)：计算智能 UAV 路径规划综述，替换为 2020、2024 年综述。",
]


KEPT_CLASSIC_REFERENCES = [
    "Hart et al. (1968)：A* 经典源文献。",
    "Stentz (1994)：D* 经典源文献。",
    "Koenig & Likhachev (2002)：D* Lite 经典源文献，原文无 DOI，保留 AAAI 页面 URL。",
    "Koenig et al. (2004)：LPA* 经典源文献。",
    "Karaman & Frazzoli (2011)：RRT* 和最优采样规划经典源文献。",
]


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(color)
    run_properties.append(underline)
    run.append(run_properties)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def replace_text_in_runs(paragraph, old: str, new: str) -> bool:
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True

    if old not in paragraph.text:
        return False

    # 兜底路径只用于引用组跨多个 run 的情况，会保留段落级样式但丢失段内局部格式。
    paragraph.text = paragraph.text.replace(old, new)
    return True


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)


def prepare_references_section(document: Document):
    heading_index = None
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip() == "References":
            heading_index = index
            break
    if heading_index is None:
        document.add_page_break()
        heading = document.add_heading("References", level=1)
    else:
        heading = document.paragraphs[heading_index]
        for paragraph in list(document.paragraphs[heading_index + 1 :]):
            remove_paragraph(paragraph)
    return heading


def add_reference_paragraph(document: Document, reference: Reference) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.5)
    paragraph.paragraph_format.space_after = 0
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.add_run(reference.text)
    if reference.link:
        paragraph.add_run(" ")
        add_hyperlink(paragraph, reference.link, reference.link)


def write_report(replacement_results: list[tuple[str, bool]]) -> None:
    eswa_count = sum(1 for reference in REFERENCES if reference.venue == "Expert Systems with Applications")
    old_count = sum(1 for reference in REFERENCES if reference.year < 2020)
    doi_count = sum(1 for reference in REFERENCES if reference.link and reference.link.startswith("https://doi.org/"))

    lines = [
        "# 初稿5 引用重构审计",
        "",
        "## 处理原则",
        "",
        "本次按用户规则执行引用重构：正文和参考文献表优先保留 2020 年及之后的文献；只有 A*、D*、D* Lite、LPA*、RRT* 这类真正承担算法谱系依据的经典文献继续保留。旧的应用型、综述型或仅背景型文献均移除或替换。",
        "",
        "## 输出统计",
        "",
        f"- 重构后参考文献总数：{len(REFERENCES)}",
        f"- 2020 年及之后文献数：{len(REFERENCES) - old_count}",
        f"- 保留的 2020 年前经典源文献数：{old_count}",
        f"- Expert Systems with Applications 文献数：{eswa_count}",
        f"- DOI 超链接文献数：{doi_count}",
        f"- 无 DOI 但保留 URL 的文献数：{len(REFERENCES) - doi_count}",
        "",
        "## 正文引用替换结果",
        "",
    ]
    for note, ok in replacement_results:
        status = "完成" if ok else "未找到原引用组"
        lines.append(f"- {status}：{note}")

    lines.extend(["", "## 保留的经典旧文献", ""])
    lines.extend(f"- {item}" for item in KEPT_CLASSIC_REFERENCES)

    lines.extend(["", "## 移除或替换的旧文献", ""])
    lines.extend(f"- {item}" for item in REMOVED_REFERENCES)

    lines.extend(["", "## 新增替换文献", ""])
    for reference in REFERENCES:
        if reference.key in {"Bertolotto2020", "Cichocinski2021", "Ghambari2024", "Mishra2024", "Simplicio2025", "Yilmaz2024"}:
            link = reference.link or "无 DOI"
            lines.append(f"- {reference.text} {link}")

    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    REPORT_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    if not INPUT_DOCX.exists():
        raise FileNotFoundError(f"找不到输入文件：{INPUT_DOCX}")

    document = Document(str(INPUT_DOCX))

    replacement_results: list[tuple[str, bool]] = []
    for old, new, note in CITATION_REPLACEMENTS:
        replaced = False
        for paragraph in document.paragraphs:
            if replace_text_in_runs(paragraph, old, new):
                replaced = True
                break
        replacement_results.append((note, replaced))

    prepare_references_section(document)
    for reference in REFERENCES:
        add_reference_paragraph(document, reference)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(OUTPUT_DOCX))
    write_report(replacement_results)

    failed = [note for note, ok in replacement_results if not ok]
    print(f"已生成 Word 文件：{OUTPUT_DOCX}")
    print(f"已生成审计报告：{REPORT_PATH}")
    print(f"正文引用替换完成数：{len(replacement_results) - len(failed)} / {len(replacement_results)}")
    print(f"参考文献条目数：{len(REFERENCES)}")
    print(f"ESWA 文献条目数：{sum(1 for item in REFERENCES if item.venue == 'Expert Systems with Applications')}")
    if failed:
        print("以下引用组未替换：")
        for item in failed:
            print(item)


if __name__ == "__main__":
    main()
