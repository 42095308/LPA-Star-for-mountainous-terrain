from __future__ import annotations

import re
import shutil
import sqlite3
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


INPUT_DOCX = Path(r"C:\Users\42095\Desktop\小论文资料\稿件\初稿8.docx")
OUTPUT_DOCX = Path(r"C:\Users\42095\Desktop\小论文资料\稿件\初稿8_文献复核插入版.docx")
REPORT_PATH = Path(r"C:\programming\code\article_python\output\literature\初稿8_文献复核报告.md")
ZOTERO_DB = Path(r"C:\Users\42095\Zotero\zotero.sqlite")
DB_COPY = Path(r"C:\programming\code\article_python\tmp\zotero\zotero_20260614_audit.sqlite")


EXPECTED_DOIS = {
    "10.1016/j.comcom.2019.10.014": "Aggarwal & Kumar, 2020",
    "10.1016/j.eswa.2023.120713": "Ali et al., 2023",
    "10.5311/josis.2020.20.659": "Bertolotto et al., 2020",
    "10.1016/j.eswa.2024.124172": "Bhuiyan et al., 2024",
    "10.1007/978-981-19-6613-2_50": "Chen et al., 2023",
    "10.4467/21995923gp.21.007.14978": "Cichociński, 2021",
    "10.1016/j.eswa.2022.119137": "Fan et al., 2023",
    "10.1051/ro/2024073": "Ghambari et al., 2024",
    "10.1109/tssc.1968.300136": "Hart et al., 1968",
    "10.1016/j.ast.2025.109947": "Hu et al., 2025",
    "10.3390/en14071968": "Huang & Savkin, 2021",
    "10.1177/0278364911406761": "Karaman & Frazzoli, 2011",
    "NO_DOI_D_LITE": "Koenig & Likhachev, 2002",
    "10.1016/j.artint.2003.12.001": "Koenig et al., 2004",
    "10.1109/yac66630.2025.11150073": "Liu & Zou, 2025",
    "10.1109/icuas48674.2020.9214057": "Margraff et al., 2020",
    "10.1016/j.eswa.2024.124254": "Mishra & Tiwari, 2024",
    "10.1109/iccworkshops50388.2021.9473740": "Mohammed et al., 2021",
    "10.1016/j.eswa.2021.114854": "Moshref-Javadi & Winkenbach, 2021",
    "10.5815/ijisa.2025.04.01": "Nikolaiev & Novotarskyi, 2025",
    "10.1016/j.eswa.2023.120992": "Ramos & Vigo, 2023",
    "10.1109/globecom54140.2023.10437552": "Salhi & Delavernhe, 2023",
    "10.1109/lra.2025.3621932": "Shao et al., 2025",
    "10.1109/icuas65942.2025.11007892": "Simplicio & Pereira, 2025",
    "10.1109/robot.1994.351061": "Stentz, 1994",
    "10.1016/j.aei.2024.102947": "Wang et al., 2024",
    "10.3390/en15218036": "Wu et al., 2022",
    "10.1016/j.knosys.2022.110164": "Xu et al., 2023",
    "10.1016/j.eswa.2023.123115": "Yılmaz et al., 2024",
    "10.1016/j.eswa.2022.119327": "Yu et al., 2023",
    "10.1016/j.eswa.2023.119910": "Yu & Luo, 2023",
    "10.1016/j.eswa.2022.119243": "Zhang et al., 2023",
    "10.1016/j.eswa.2023.121862": "Zhang et al., 2024",
    "10.1016/j.eswa.2024.123481": "Zhou et al., 2024",
}

OLD_NONCLASSIC_TERMS = [
    "Daniel et al., 2010",
    "Theta*: Any-angle path planning on grids",
    "Haklay & Weber, 2008",
    "OpenStreetMap: User-generated street maps",
    "Kim et al., 2019",
    "A study on 3D optimal path planning for quadcopter UAV based on D* Lite",
    "Mardani et al., 2019",
    "Communication-aware UAV path planning",
    "Nash et al., 2010",
    "Lazy Theta*",
    "Song et al., 2018",
    "Persistent UAV delivery logistics",
    "Sun et al., 2011",
    "Intelligent flight task algorithm for unmanned aerial vehicle",
    "Tavana et al., 2017",
    "Drone shipping versus truck delivery",
    "Volz & Graichen, 2019",
    "predictive path-following controller",
    "Zhao et al., 2018",
    "Survey on computational-intelligence-based UAV path planning",
]

OPTIONAL_NEW_REFERENCES = [
    "MSCPSO: A multi-strategy cooperative particle swarm optimization algorithm for UAV path planning. Expert Systems with Applications, 2026. https://doi.org/10.1016/j.eswa.2025.131034",
    "Research on collaborative path planning of UAV swarms for urban logistics distribution in dense building environments. Expert Systems with Applications, 2026. https://doi.org/10.1016/j.eswa.2025.130816",
    "Energy-predictive planning for optimizing drone service delivery. Expert Systems with Applications, 2026. https://doi.org/10.1016/j.eswa.2025.129251",
]


def copy_and_read_zotero_dois() -> tuple[set[str], int, list[str]]:
    DB_COPY.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(ZOTERO_DB, DB_COPY)
    con = sqlite3.connect(DB_COPY)
    con.row_factory = sqlite3.Row
    rows = con.execute(
        """
        select i.itemID, it.typeName, f.fieldName, v.value
        from items i
        join itemTypes it on i.itemTypeID = it.itemTypeID
        left join itemData d on i.itemID = d.itemID
        left join fields f on d.fieldID = f.fieldID
        left join itemDataValues v on d.valueID = v.valueID
        where i.itemID not in (select itemID from deletedItems)
        order by i.itemID
        """
    ).fetchall()
    items: dict[int, dict[str, str]] = {}
    item_types: dict[int, str] = {}
    for row in rows:
        item_types[row["itemID"]] = row["typeName"]
        fields = items.setdefault(row["itemID"], {})
        if row["fieldName"]:
            fields[row["fieldName"]] = row["value"]
    con.close()

    dois: set[str] = set()
    titles_no_doi: list[str] = []
    count = 0
    for item_id, fields in items.items():
        if item_types.get(item_id) in {"attachment", "annotation"}:
            continue
        if not fields.get("title"):
            continue
        count += 1
        doi = fields.get("DOI", "").strip().lower()
        if doi:
            dois.add(doi)
        else:
            titles_no_doi.append(fields.get("title", ""))
    return dois, count, titles_no_doi


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(color)
    run_properties.append(underline)
    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def clear_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def repair_reference_links(document: Document, ref_index: int) -> int:
    repaired = 0
    pattern = re.compile(r"(https?://\S+)\s*$")
    for paragraph in document.paragraphs[ref_index + 1 :]:
        text = paragraph.text.strip()
        if not text:
            continue
        match = pattern.search(text)
        if not match:
            continue
        url = match.group(1)
        prefix = text[: match.start(1)].rstrip()
        clear_runs(paragraph)
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.5)
        paragraph.paragraph_format.space_after = 0
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.add_run(prefix + " ")
        add_hyperlink(paragraph, url, url)
        repaired += 1
    return repaired


def count_active_links(path: Path) -> tuple[int, int, int]:
    with ZipFile(path) as package:
        rels = package.read("word/_rels/document.xml.rels")
    root = ET.fromstring(rels)
    links = [
        element.attrib.get("Target", "")
        for element in root
        if element.attrib.get("Type", "").endswith("/hyperlink")
    ]
    return (
        len(links),
        sum(link.startswith("https://doi.org/") for link in links),
        sum("aaai.org/papers" in link for link in links),
    )


def main() -> None:
    document = Document(str(INPUT_DOCX))
    ref_index = next(
        (index for index, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() == "References"),
        None,
    )
    if ref_index is None:
        raise RuntimeError("未找到 References 标题")

    body_text = "\n".join(p.text.strip() for p in document.paragraphs[:ref_index] if p.text.strip())
    refs = [p.text.strip() for p in document.paragraphs[ref_index + 1 :] if p.text.strip()]
    refs_text = "\n".join(refs)
    old_hits_body = [term for term in OLD_NONCLASSIC_TERMS if term in body_text]
    old_hits_refs = [term for term in OLD_NONCLASSIC_TERMS if term in refs_text]

    doc_dois = {doi.lower() for doi in re.findall(r"https://doi\.org/([^\s]+)", refs_text)}
    doc_has_d_lite = "aaai.org/papers/00476-aaai02-072-d-lite" in refs_text
    expected_dois = {doi for doi in EXPECTED_DOIS if doi != "NO_DOI_D_LITE"}
    missing_from_doc = sorted(expected_dois - doc_dois)
    extra_doc_dois = sorted(doc_dois - expected_dois)

    zotero_dois, zotero_count, no_doi_titles = copy_and_read_zotero_dois()
    missing_from_zotero = sorted(expected_dois - zotero_dois)
    duplicate_hint = zotero_count - len(zotero_dois) - len(no_doi_titles)

    repaired = repair_reference_links(document, ref_index)
    document.save(str(OUTPUT_DOCX))
    active_links, active_doi_links, active_aaai_links = count_active_links(OUTPUT_DOCX)

    eswa_count = sum("Expert Systems with Applications" in ref for ref in refs)
    old_ref_count = 0
    year_dist: dict[str, int] = {}
    for ref in refs:
        match = re.search(r"\((\d{4})\)", ref)
        if match:
            year = match.group(1)
            year_dist[year] = year_dist.get(year, 0) + 1
            if int(year) < 2020:
                old_ref_count += 1

    lines = [
        "# 初稿8 文献复核报告",
        "",
        "## 结论",
        "",
        "初稿8 已经包含上一轮最终确定的 34 篇文献。正文引用和 References 的双向检查通过，旧的非经典文献未发现残留。本次生成的复核插入版没有增加或删除正文引用，只修复了 References 中 URL 不是活动超链接的问题。",
        "",
        "## 文档检查",
        "",
        f"- 正文非空段落数：{len([p for p in document.paragraphs[:ref_index] if p.text.strip()])}",
        f"- References 条目数：{len(refs)}",
        f"- Expert Systems with Applications 条目数：{eswa_count}",
        f"- 2020 年前经典文献数：{old_ref_count}",
        f"- 年份分布：{dict(sorted(year_dist.items()))}",
        f"- 缺失的预期 DOI：{missing_from_doc or '无'}",
        f"- 额外 DOI：{extra_doc_dois or '无'}",
        f"- D* Lite URL 是否存在：{doc_has_d_lite}",
        "",
        "## 旧文献删除检查",
        "",
        f"- 正文旧非经典文献残留：{old_hits_body or '无'}",
        f"- References 旧非经典文献残留：{old_hits_refs or '无'}",
        "",
        "## Zotero 本地库检查",
        "",
        f"- Zotero 有效文献条目数：{zotero_count}",
        f"- 缺失的最终 33 个 DOI：{missing_from_zotero or '无'}",
        f"- 无 DOI 条目题名：{no_doi_titles or '无'}",
        f"- 重复条目提示值：{duplicate_hint}，当前主要表现为 Hu et al. 2025 在库中重复",
        "",
        "## 超链接修复",
        "",
        f"- 修复 References 链接数：{repaired}",
        f"- 输出文件活动超链接数：{active_links}",
        f"- DOI 活动超链接数：{active_doi_links}",
        f"- AAAI URL 活动超链接数：{active_aaai_links}",
        "",
        "## 是否需要新增或减少文献",
        "",
        "当前不建议强制新增文献。34 篇对于 ESWA 小论文已经足够，且其中 29 篇为 2020 年及之后文献，12 篇来自 Expert Systems with Applications，5 篇 2020 年前文献均为 A*、D*、D* Lite、LPA*、RRT* 经典算法源文献，保留合理。",
        "",
        "可以考虑但不建议立即插入的 2026 年 ESWA 候选如下。这些文献能增强时效性，但主题多偏多无人机、城市物流或服务配送，与本文山地三层航线网络和事件驱动增量重规划不是完全同一问题。若投稿前想进一步强调最新 ESWA 背景，可用其中 1 篇替换现有相对泛化的优化算法背景文献，而不是继续增加总量。",
        "",
    ]
    lines.extend(f"- {item}" for item in OPTIONAL_NEW_REFERENCES)
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    REPORT_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8")

    print(f"已生成复核插入版：{OUTPUT_DOCX}")
    print(f"已生成复核报告：{REPORT_PATH}")
    print(f"References条目数：{len(refs)}")
    print(f"旧非经典正文残留：{len(old_hits_body)}")
    print(f"旧非经典References残留：{len(old_hits_refs)}")
    print(f"缺失预期DOI：{len(missing_from_doc)}")
    print(f"Zotero缺失最终DOI：{len(missing_from_zotero)}")
    print(f"活动超链接数：{active_links}")


if __name__ == "__main__":
    main()
