from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


INPUT_DOCX = Path(r"C:\Users\42095\Desktop\初稿6-更正意见版_文献插入版.docx")
OUTPUT_DOCX = Path(r"C:\Users\42095\Desktop\初稿6-更正意见版_文献插入版_超链接修正版.docx")


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(color)
    run_properties.append(underline)
    run.append(run_properties)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def clear_paragraph_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def main() -> None:
    document = Document(str(INPUT_DOCX))
    ref_index = None
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip() == "References":
            ref_index = index
            break
    if ref_index is None:
        raise RuntimeError("未找到 References 标题")

    repaired = 0
    pattern = re.compile(r"(https?://\S+)\s*$")
    for paragraph in document.paragraphs[ref_index + 1 :]:
        text = paragraph.text.strip()
        if not text:
            continue
        match = pattern.search(text)
        if not match:
            continue
        url = match.group(1)
        prefix = text[: match.start(1)].rstrip()
        clear_paragraph_runs(paragraph)
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.5)
        paragraph.paragraph_format.space_after = 0
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.add_run(prefix + " ")
        add_hyperlink(paragraph, url, url)
        repaired += 1

    document.save(str(OUTPUT_DOCX))
    print(f"已生成超链接修正版：{OUTPUT_DOCX}")
    print(f"修复超链接数量：{repaired}")


if __name__ == "__main__":
    main()
